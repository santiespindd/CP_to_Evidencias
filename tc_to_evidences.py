import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from shutil import copyfile

def generar_evidencias():
    archivo_excel = archivo_entry.get()
    carpeta_destino = destino_entry.get()
    nombre_archivo = nombre_entry.get()
    
    # Verificar si se seleccionaron el archivo y la carpeta de destino
    if archivo_excel == "" or carpeta_destino == "" or nombre_archivo == "":
        messagebox.showerror("Error", "Por favor, completa todos los campos.")
        return
    
    try:
        # Leer el archivo Excel
        df = pd.read_excel(archivo_excel)
    except Exception as e:
        messagebox.showerror("Error", "No se pudo leer el archivo Excel. Asegúrate de seleccionar un archivo válido.")
        return
    
    # Verificar el formato del archivo Excel
    if "Caso de Prueba" not in df.columns or "Pasos" not in df.columns or "Resultado Esperado" not in df.columns:
        messagebox.showerror("Error", "El formato del Excel no es el correcto. Debe contener las columnas 'Caso de Prueba', 'Pasos' y 'Resultado Esperado'.")
        return
    
    # Crear un documento Word
    doc = Document()
    
    # Agregar título al documento
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("Evidencias")
    title_font = title_run.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Subtítulos
    subtitles = ['Tester:', 'Ambiente:', 'Fecha:']
    for subtitle in subtitles:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.bold = True
    
    doc.add_heading(f'Caso de Prueba ejecutados', level=1)
    
    paso = 0
    
    # Iterar a través de los casos de Prueba
    for index, row in df.iterrows():
        # Agregar sección para cada caso de Prueba
        if pd.notna(row["Caso de Prueba"]) and row["Caso de Prueba"] != '':
            doc.add_heading(f'{row["Caso de Prueba"]}', level=2)
            paso = 1
        doc.add_paragraph(f'Pasos {paso}: {row["Pasos"]}')
        paso += 1
        if pd.notna(row["Resultado Esperado"]) and row["Resultado Esperado"] != '':
            doc.add_paragraph(f'Resultado Esperado: {row["Resultado Esperado"]}')
    
    # Guardar el documento Word en la carpeta de destino con el nombre especificado
    ruta_archivo = f"{carpeta_destino}/{nombre_archivo}.docx"
    doc.save(ruta_archivo)
    
    messagebox.showinfo("Éxito", "Evidencias generadas correctamente.")

def seleccionar_archivo():
    archivo = filedialog.askopenfilename(title="Seleccionar archivo Excel")
    archivo_entry.delete(0, tk.END)
    archivo_entry.insert(0, archivo)

def seleccionar_destino():
    destino = filedialog.askdirectory(title="Seleccionar carpeta de destino")
    destino_entry.delete(0, tk.END)
    destino_entry.insert(0, destino)

# Crear la ventana principal
root = tk.Tk()
root.title("Generar archivo de evidencias")

# Etiqueta y entrada para el archivo Excel
archivo_label = tk.Label(root, text="Archivo Excel:")
archivo_label.grid(row=0, column=0, padx=10, pady=5)
archivo_entry = tk.Entry(root, width=50)
archivo_entry.grid(row=0, column=1, padx=10, pady=5)
btn_seleccionar_archivo = tk.Button(root, text="Seleccionar", command=seleccionar_archivo)
btn_seleccionar_archivo.grid(row=0, column=2, padx=10, pady=5)

# Etiqueta y entrada para la carpeta de destino
destino_label = tk.Label(root, text="Carpeta de destino:")
destino_label.grid(row=1, column=0, padx=10, pady=5)
destino_entry = tk.Entry(root, width=50)
destino_entry.grid(row=1, column=1, padx=10, pady=5)
btn_seleccionar_destino = tk.Button(root, text="Seleccionar", command=seleccionar_destino)
btn_seleccionar_destino.grid(row=1, column=2, padx=10, pady=5)

# Etiqueta y entrada para el nombre del archivo
nombre_label = tk.Label(root, text="Nombre del archivo:")
nombre_label.grid(row=2, column=0, padx=10, pady=5)
nombre_entry = tk.Entry(root, width=50)
nombre_entry.grid(row=2, column=1, padx=10, pady=5)

# Botón para generar evidencias
btn_generar = tk.Button(root, text="Generar Evidencias", command=generar_evidencias)
btn_generar.grid(row=3, column=1, padx=10, pady=20)

# Ejecutar la interfaz
root.mainloop()
